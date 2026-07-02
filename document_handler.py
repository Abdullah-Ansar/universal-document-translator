from docx import Document


def read_document(file_path):

    doc = Document(file_path)

    data = []

    for para in doc.paragraphs:

        if para.text.strip():
            data.append(para.text)

    for table in doc.tables:

        for row in table.rows:

            for cell in row.cells:

                if cell.text.strip():

                    data.append(cell.text)

    return data