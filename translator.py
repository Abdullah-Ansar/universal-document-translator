from docx import Document
import re

def is_hindi(text):
    return bool(re.search(r'[\u0900-\u097F]', text))


def read_docx(file_path):
    doc = Document(file_path)

    paragraphs = []

    # Read paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text)

    # Read tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text)

    return paragraphs