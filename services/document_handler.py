from docx import Document


class DocumentHandler:

    def __init__(self, file_path):
        self.document = Document(file_path)

    def get_paragraphs(self):
        return self.document.paragraphs

    def get_tables(self):
        return self.document.tables

    def save(self, output_path):
        self.document.save(output_path)